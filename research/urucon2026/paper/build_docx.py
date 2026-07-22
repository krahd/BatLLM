#!/usr/bin/env python3
"""Build the editable British-English DOCX version of the URUCON paper.

The PDF remains the authoritative IEEE-formatted submission. This script uses
Pandoc for LaTeX mathematics and citation conversion, then applies a stable
single-column academic layout with python-docx.
"""

from __future__ import annotations

import argparse
from pathlib import Path
import shutil
import subprocess
import tempfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def _run_font(run, size: float | None = None, bold=None, italic=None) -> None:
    name = "Times New Roman"
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "en-GB")
    lang.set(qn("w:eastAsia"), "en-GB")
    lang.set(qn("w:bidi"), "en-GB")


def _style_font(style, size: float, bold=None, italic=None) -> None:
    name = "Times New Roman"
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "en-GB")


def _insert_after(paragraph, new_paragraph) -> None:
    paragraph._p.addnext(new_paragraph._p)


def _page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    _run_font(run, 9)


def _repeat_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    trpr.append(node)


def _prevent_split(row) -> None:
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def _cell_margins(cell) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    margins = tcpr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tcpr.append(margins)
    for name, value in (("top", 70), ("start", 80), ("bottom", 70), ("end", 80)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _locate_ieee_csl() -> Path:
    candidates = [
        Path(
            "/usr/share/texlive/texmf-dist/tex/latex/"
            "citation-style-language/styles/ieee.csl"
        ),
        Path("/usr/share/citation-style-language/styles/ieee.csl"),
        Path("/usr/share/pandoc/data/csl/ieee.csl"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    raise FileNotFoundError(
        "IEEE CSL file not found; install TeX Live citation-style-language styles."
    )


def _preprocess_latex(source: Path, destination: Path) -> None:
    text = source.read_text(encoding="utf-8")
    text = text.replace(
        "\\section*{Acknowledgment}", "\\section*{Acknowledgement}"
    )
    start = text.find("\\begin{thebibliography}")
    end = text.find("\\end{thebibliography}")
    if start >= 0 and end >= 0:
        end += len("\\end{thebibliography}")
        text = text[:start] + "\n" + text[end:]
    destination.write_text(text, encoding="utf-8")


def style_docx(raw_path: Path, output_path: Path) -> None:
    doc = Document(raw_path)
    doc.core_properties.title = (
        "From Prompt to State: Verifiable Grounding and Operative Replay "
        "for LLM-Mediated Control"
    )
    doc.core_properties.subject = "URUCON 2026 Computing Track paper"
    doc.core_properties.author = "Tomas Laurenzo"
    doc.core_properties.language = "en-GB"
    doc.core_properties.keywords = (
        "large language models, provenance, replay, grounding, verification, "
        "human-AI interaction"
    )

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.header_distance = Cm(0.9)
        section.footer_distance = Cm(0.9)
        footer = section.footer.paragraphs[0]
        footer.clear()
        _page_number(footer)

    sizes = {
        "Normal": 10.5,
        "Body Text": 10.5,
        "First Paragraph": 10.5,
        "Abstract": 9.5,
        "Abstract Title": 10.5,
        "Title": 16,
        "Heading 1": 12,
        "Heading 2": 10.5,
        "Table Caption": 9.5,
        "Bibliography": 9,
    }
    for name, size in sizes.items():
        if name in doc.styles:
            _style_font(doc.styles[name], size)

    for name in ("Normal", "Body Text", "First Paragraph"):
        if name in doc.styles:
            fmt = doc.styles[name].paragraph_format
            fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            fmt.line_spacing = 1.02
            fmt.space_after = Pt(4)
            fmt.widow_control = True
    if "Title" in doc.styles:
        fmt = doc.styles["Title"].paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.space_after = Pt(8)
    if "Heading 1" in doc.styles:
        fmt = doc.styles["Heading 1"].paragraph_format
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(4)
        fmt.keep_with_next = True
    if "Heading 2" in doc.styles:
        fmt = doc.styles["Heading 2"].paragraph_format
        fmt.space_before = Pt(8)
        fmt.space_after = Pt(3)
        fmt.keep_with_next = True
        doc.styles["Heading 2"].font.italic = True
    if "Abstract" in doc.styles:
        fmt = doc.styles["Abstract"].paragraph_format
        fmt.left_indent = Cm(0.7)
        fmt.right_indent = Cm(0.7)
        fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.line_spacing = 1.0
    if "Abstract Title" in doc.styles:
        fmt = doc.styles["Abstract Title"].paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.keep_with_next = True
        doc.styles["Abstract Title"].font.bold = True
    if "Table Caption" in doc.styles:
        fmt = doc.styles["Table Caption"].paragraph_format
        fmt.space_before = Pt(6)
        fmt.space_after = Pt(3)
        fmt.keep_with_next = True
        doc.styles["Table Caption"].font.italic = True
    if "Bibliography" in doc.styles:
        fmt = doc.styles["Bibliography"].paragraph_format
        fmt.left_indent = Cm(0.8)
        fmt.first_line_indent = Cm(-0.8)
        fmt.line_spacing = 1.0
        fmt.space_after = Pt(2)

    title = doc.paragraphs[0]
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_font(author.add_run("Tomas Laurenzo"), 11)
    _insert_after(title, author)
    affiliation = doc.add_paragraph()
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliation.paragraph_format.space_after = Pt(9)
    _run_font(
        affiliation.add_run(
            "Department of Critical Media Practices\n"
            "University of Colorado Boulder\n"
            "Boulder, Colorado, USA\n"
            "tomas@laurenzo.net"
        ),
        9.5,
    )
    _insert_after(author, affiliation)

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(
            "large language models, provenance, replay"
        ):
            words = paragraph.text.strip()
            paragraph.clear()
            paragraph.paragraph_format.left_indent = Cm(0.7)
            paragraph.paragraph_format.right_indent = Cm(0.7)
            paragraph.paragraph_format.space_after = Pt(8)
            _run_font(paragraph.add_run("Keywords—"), 9.5, bold=True)
            _run_font(paragraph.add_run(words), 9.5, italic=True)
            break

    caption_number = 0
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Table Caption":
            caption_number += 1
            original = paragraph.text.strip()
            paragraph.clear()
            _run_font(
                paragraph.add_run(f"Table {caption_number}. "), 9.5, bold=True
            )
            _run_font(paragraph.add_run(original), 9.5, italic=True)

    for paragraph in doc.paragraphs:
        style = paragraph.style.name
        for run in paragraph.runs:
            if style == "Title":
                _run_font(run, 16, bold=True)
            elif style == "Heading 1":
                _run_font(run, 12, bold=True)
            elif style == "Heading 2":
                _run_font(run, 10.5, bold=True, italic=True)
            elif style == "Abstract Title":
                _run_font(run, 10.5, bold=True)
            elif style == "Abstract":
                _run_font(run, 9.5)
            elif style == "Bibliography":
                _run_font(run, 9)
            elif style != "Table Caption":
                _run_font(run, 10.5)
            if "\u00a0" in run.text:
                run.text = run.text.replace("\u00a0", " ")

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        if table.rows:
            _repeat_header(table.rows[0])
        for row_index, row in enumerate(table.rows):
            _prevent_split(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                _cell_margins(cell)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    if row_index < len(table.rows) - 1:
                        paragraph.paragraph_format.keep_with_next = True
                    for run in paragraph.runs:
                        _run_font(run, 9.2, bold=(row_index == 0))

    settings = doc.settings._element
    language = settings.find(qn("w:themeFontLang"))
    if language is None:
        language = OxmlElement("w:themeFontLang")
        settings.append(language)
    language.set(qn("w:val"), "en-GB")
    language.set(qn("w:eastAsia"), "en-GB")

    doc.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--source", type=Path, default=Path(__file__).with_name("main.tex")
    )
    parser.add_argument(
        "--bibliography",
        type=Path,
        default=Path(__file__).with_name("references.bib"),
    )
    parser.add_argument(
        "--output", type=Path, default=Path(__file__).with_name("main.docx")
    )
    args = parser.parse_args()

    if shutil.which("pandoc") is None:
        raise RuntimeError("pandoc is required")
    args.output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="batllm-docx-") as tmp:
        tmpdir = Path(tmp)
        prepared = tmpdir / "paper.tex"
        raw = tmpdir / "paper.docx"
        _preprocess_latex(args.source, prepared)
        subprocess.run(
            [
                "pandoc",
                str(prepared),
                "--from=latex",
                "--to=docx",
                "--citeproc",
                f"--bibliography={args.bibliography}",
                f"--csl={_locate_ieee_csl()}",
                "--metadata=reference-section-title:References",
                f"--output={raw}",
            ],
            cwd=args.source.parent,
            check=True,
        )
        style_docx(raw, args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
